from flask import Blueprint, make_response, send_file
from flask_login import login_required
from app.models import Student, Teacher, FeePayment
from fpdf import FPDF
from docx import Document
import io
import csv

export_bp = Blueprint('export', __name__)

@export_bp.route('/students/pdf')
@login_required
def students_pdf():
    students = Student.query.all()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt="Student Report", ln=1, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    pdf.cell(30, 10, "Reg No", 1)
    pdf.cell(50, 10, "Name", 1)
    pdf.cell(50, 10, "Father Name", 1)
    pdf.cell(30, 10, "Status", 1)
    pdf.ln()
    
    for s in students:
        pdf.cell(30, 10, str(s.reg_no), 1)
        pdf.cell(50, 10, str(s.name), 1)
        pdf.cell(50, 10, str(s.father_name), 1)
        pdf.cell(30, 10, str(s.status), 1)
        pdf.ln()

    output = io.BytesIO(pdf.output(dest='S').encode('latin-1'))
    output.seek(0)
    
    return send_file(output, download_name='students.pdf', as_attachment=True, mimetype='application/pdf')

@export_bp.route('/students/word')
@login_required
def students_word():
    students = Student.query.all()
    doc = Document()
    doc.add_heading('Student Report', 0)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reg No'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Father Name'
    hdr_cells[3].text = 'Status'
    
    for s in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(s.reg_no)
        row_cells[1].text = str(s.name)
        row_cells[2].text = str(s.father_name)
        row_cells[3].text = str(s.status)
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    return send_file(f, download_name='students.docx', as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@export_bp.route('/teachers/pdf')
@login_required
def teachers_pdf():
    teachers = Teacher.query.all()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt="Teacher Report", ln=1, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    pdf.cell(30, 10, "Emp ID", 1)
    pdf.cell(50, 10, "Name", 1)
    pdf.cell(50, 10, "Designation", 1)
    pdf.cell(40, 10, "Contact", 1)
    pdf.ln()
    
    for t in teachers:
        pdf.cell(30, 10, str(t.employee_id), 1)
        pdf.cell(50, 10, str(t.name), 1)
        pdf.cell(50, 10, str(t.designation), 1)
        pdf.cell(40, 10, str(t.contact), 1)
        pdf.ln()

    output = io.BytesIO(pdf.output(dest='S').encode('latin-1'))
    output.seek(0)
    
    return send_file(output, download_name='teachers.pdf', as_attachment=True, mimetype='application/pdf')

@export_bp.route('/teachers/word')
@login_required
def teachers_word():
    teachers = Teacher.query.all()
    doc = Document()
    doc.add_heading('Teacher Report', 0)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Emp ID'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Designation'
    hdr_cells[3].text = 'Contact'
    
    for t in teachers:
        row_cells = table.add_row().cells
        row_cells[0].text = str(t.employee_id)
        row_cells[1].text = str(t.name)
        row_cells[2].text = str(t.designation)
        row_cells[3].text = str(t.contact)
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    return send_file(f, download_name='teachers.docx', as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
